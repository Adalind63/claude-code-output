# -*- coding: utf-8 -*-
"""生成模型使用指南 Word 文档"""
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

style = doc.styles['Normal']
font = style.font
font.name = 'Consolas'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def add_title(text):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(text)
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers), style='Light Grid Accent 1')
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(10)
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()

def add_code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x2d, 0x2d, 0x2d)

# ============ 封面 ============
doc.add_paragraph()
doc.add_paragraph()
add_title('BERT 立场分类模型\n使用指南')
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('DistilBERT · Tough / Cooperate / Neutral')
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run(datetime.now().strftime('%Y-%m-%d'))
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
doc.add_page_break()

# ============ 1. 文件结构 ============
add_heading('1. 文件路径总览', level=1)

add_table(
    ['类型', '路径', '说明'],
    [
        ['模型目录', r'D:\项目流程\bert_stance_model', '可直接加载的最终模型'],
        ['模型权重', r'D:\项目流程\bert_stance_model\model.safetensors', '255 MB，DistilBERT 微调权重'],
        ['模型配置', r'D:\项目流程\bert_stance_model\config.json', '模型结构定义'],
        ['分词器', r'D:\项目流程\bert_stance_model\tokenizer.json', '695 KB，WordPiece 词表'],
        ['分词器配置', r'D:\项目流程\bert_stance_model\tokenizer_config.json', '分词参数'],
        ['训练报告', r'D:\项目流程\BERT立场微调训练报告.docx', '完整训练报告（含图表）'],
        ['使用指南', r'D:\项目流程\BERT模型使用指南.docx', '本文件'],
        ['训练数据', r'D:\项目流程\train.csv', '训练集'],
        ['验证数据', r'D:\项目流程\val.csv', '验证集'],
        ['测试数据', r'D:\项目流程\test.csv', '测试集'],
        ['训练脚本', r'D:\项目流程\bert_stance_finetune.py', '原始训练脚本'],
        ['续跑脚本', r'D:\项目流程\resume_training.py', 'Checkpoint 续跑脚本'],
        ['报告生成', r'D:\项目流程\generate_report.py', 'Word 报告生成脚本'],
        ['全部 Checkpoints', r'D:\项目流程\bert_stance_model\checkpoints', '各轮次存档，可清理'],
    ]
)

# ============ 2. 环境要求 ============
add_heading('2. 环境要求', level=1)

add_table(
    ['依赖库', '版本', '安装命令'],
    [
        ['torch', '≥ 2.0', 'pip install torch'],
        ['transformers', '≥ 4.30', 'pip install transformers'],
        ['datasets', '≥ 2.0', 'pip install datasets'],
        ['scikit-learn', '≥ 1.0', 'pip install scikit-learn'],
        ['pandas', '≥ 1.5', 'pip install pandas'],
        ['numpy', '≥ 1.24', 'pip install numpy'],
    ]
)

doc.add_paragraph('一键安装：')
add_code('pip install torch transformers datasets scikit-learn pandas numpy')

# ============ 3. 快速使用 ============
add_heading('3. 快速使用', level=1)

add_heading('3.1 基础推理（单条文本）', level=2)
doc.add_paragraph('加载模型并对一条文本进行立场预测：')

add_code('from transformers import DistilBertTokenizer, DistilBertForSequenceClassification')
add_code('')
add_code("# 1. 加载模型和分词器")
add_code(r'model_path = r"D:\项目流程\bert_stance_model"')
add_code('tokenizer = DistilBertTokenizer.from_pretrained(model_path)')
add_code('model = DistilBertForSequenceClassification.from_pretrained(model_path)')
add_code('')
add_code("# 2. 待预测文本")
add_code('text = "We strongly oppose China\'s unfair trade practices and will take countermeasures."')
add_code('')
add_code("# 3. Tokenize")
add_code('inputs = tokenizer(text, return_tensors="pt", truncation=True, max_length=128)')
add_code('')
add_code("# 4. 推理")
add_code('import torch')
add_code('with torch.no_grad():')
add_code('    outputs = model(**inputs)')
add_code('    logits = outputs.logits')
add_code('    pred = torch.argmax(logits, dim=1).item()')
add_code('')
add_code("# 5. 映射标签")
add_code('label_map = {0: "tough", 1: "cooperate", 2: "neutral"}')
add_code('print(f"预测立场: {label_map[pred]}")')

add_heading('3.2 批量推理（CSV 文件）', level=2)
doc.add_paragraph('从 CSV 文件读取文本，批量预测并输出结果：')

add_code('import pandas as pd')
add_code('import torch')
add_code('from transformers import DistilBertTokenizer, DistilBertForSequenceClassification')
add_code('from torch.utils.data import DataLoader, Dataset')
add_code('')
add_code('LABEL_MAP = {0: "tough", 1: "cooperate", 2: "neutral"}')
add_code(r"MODEL_PATH = r'D:\项目流程\bert_stance_model'")
add_code('')
add_code('class TextDataset(Dataset):')
add_code('    def __init__(self, texts, tokenizer, max_len=128):')
add_code('        self.encodings = tokenizer(texts, truncation=True, padding=True,')
add_code('                                   max_length=max_len, return_tensors="pt")')
add_code('    def __len__(self):')
add_code('        return len(self.encodings["input_ids"])')
add_code('    def __getitem__(self, idx):')
add_code('        return {k: v[idx] for k, v in self.encodings.items()}')
add_code('')
add_code('# 加载模型')
add_code('tokenizer = DistilBertTokenizer.from_pretrained(MODEL_PATH)')
add_code('model = DistilBertForSequenceClassification.from_pretrained(MODEL_PATH)')
add_code('model.eval()')
add_code('')
add_code('# 读取数据')
add_code(r"df = pd.read_csv(r'your_file.csv', encoding='utf-8-sig')")
add_code('texts = df["text"].astype(str).tolist()')
add_code('')
add_code('# 批量推理')
add_code('dataset = TextDataset(texts, tokenizer)')
add_code('loader = DataLoader(dataset, batch_size=32)')
add_code('')
add_code('all_preds = []')
add_code('with torch.no_grad():')
add_code('    for batch in loader:')
add_code('        logits = model(**batch).logits')
add_code('        preds = torch.argmax(logits, dim=1).tolist()')
add_code('        all_preds.extend(preds)')
add_code('')
add_code("df['stance'] = [LABEL_MAP[p] for p in all_preds]")
add_code(r"df.to_csv(r'output_with_stance.csv', index=False, encoding='utf-8-sig')")
add_code('print("Done!")')

add_heading('3.3 Pipeline 一行推理', level=2)
doc.add_paragraph('使用 Hugging Face Pipeline（最简单）：')

add_code('from transformers import pipeline')
add_code('')
add_code(r"classifier = pipeline('text-classification', model=r'D:\项目流程\bert_stance_model')")
add_code('')
add_code('result = classifier("We support deeper cooperation with China.")')
add_code('print(result)  # [{"label": "cooperate", "score": 0.92}]')

# ============ 4. 模型参数 ============
add_heading('4. 模型规格', level=1)

add_table(
    ['项目', '值'],
    [
        ['基座模型', 'distilbert-base-uncased'],
        ['任务类型', '三分类立场识别'],
        ['标签', '0=tough / 1=cooperate / 2=neutral'],
        ['最大序列长度', '128 tokens'],
        ['参数量', '66,955,011'],
        ['模型大小', '255 MB (safetensors)'],
        ['最佳 Macro F1', '0.6236 (Epoch 6)'],
        ['输入', '英文文本字符串'],
        ['输出', 'label + score（或 logits）'],
    ]
)

# ============ 5. 完整推理脚本 ============
add_heading('5. 完整推理脚本', level=1)
doc.add_paragraph('将以下代码保存为 .py 文件即可独立运行：')

add_code(r'''# -*- coding: utf-8 -*-
"""BERT 立场分类推理脚本 —— 开箱即用"""
import torch
from transformers import DistilBertTokenizer, DistilBertForSequenceClassification

MODEL_PATH = r"D:\项目流程\bert_stance_model"
LABEL_MAP = {0: "tough", 1: "cooperate", 2: "neutral"}

class StanceClassifier:
    def __init__(self, model_path=MODEL_PATH):
        self.tokenizer = DistilBertTokenizer.from_pretrained(model_path)
        self.model = DistilBertForSequenceClassification.from_pretrained(model_path)
        self.model.eval()
        if torch.cuda.is_available():
            self.model = self.model.cuda()
        self.device = "GPU" if torch.cuda.is_available() else "CPU"
        print(f"模型已加载，设备: {self.device}")

    def predict(self, text: str) -> dict:
        """预测单条文本"""
        inputs = self.tokenizer(text, return_tensors="pt",
                                truncation=True, max_length=128)
        if self.model.device.type == "cuda":
            inputs = {k: v.cuda() for k, v in inputs.items()}
        with torch.no_grad():
            logits = self.model(**inputs).logits
            probs = torch.softmax(logits, dim=1)[0]
            pred = torch.argmax(logits, dim=1).item()
        return {
            "text": text[:100] + "..." if len(text) > 100 else text,
            "label": LABEL_MAP[pred],
            "confidence": round(probs[pred].item(), 4),
            "scores": {LABEL_MAP[i]: round(probs[i].item(), 4) for i in range(3)},
        }

    def predict_batch(self, texts: list[str]) -> list[dict]:
        """批量预测"""
        return [self.predict(t) for t in texts]


if __name__ == "__main__":
    clf = StanceClassifier()

    # 示例
    samples = [
        "We must take a firm stance against unfair trade practices.",
        "China and the EU should deepen economic cooperation.",
        "The meeting was held to discuss bilateral relations.",
    ]
    for s in samples:
        result = clf.predict(s)
        print(f"\nText: {result['text']}")
        print(f"Label: {result['label']} (confidence: {result['confidence']})")
        print(f"Scores: {result['scores']}")
''')

# ============ 6. 注意事项 ============
add_heading('6. 注意事项', level=1)

notes = [
    '输入文本必须为英文。若输入中文，模型仍会输出结果，但预测不可靠。',
    '文本超过 128 tokens 会被自动截断（truncation），仅保留前 128 个 token。',
    '模型在 EU-China 政策语料上训练，领域外文本（如社交媒体、文学）性能可能下降。',
    'GPU 推理约 10-50 ms/条；CPU 推理约 100-500 ms/条。',
    '要清理训练产生的中间文件，删除 D:\\项目流程\\bert_stance_model\\checkpoints 目录（~3.5 GB）。',
    '要清理临时脚本，删除 D:\\项目流程\\resume_training.py 和 generate_report.py。',
]
for note in notes:
    p = doc.add_paragraph(note, style='List Bullet')

doc.add_paragraph()
doc.add_paragraph(f'文档生成时间：{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
doc.add_paragraph('训练最佳模型：Epoch 6 | Macro F1 = 0.6236 | Checkpoint-22164')

# ============ 保存 ============
output = r"D:\项目流程\BERT模型使用指南.docx"
doc.save(output)
print(f"使用指南已保存：{output}")
