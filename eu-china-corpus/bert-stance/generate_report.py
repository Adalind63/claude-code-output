# -*- coding: utf-8 -*-
"""生成 BERT 立场微调完整 Word 报告（含图表）"""
import json, os
from datetime import datetime
import numpy as np
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.ticker as ticker
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# ===== 配置 =====
TRAINER_STATE = r"D:\项目流程\bert_stance_model\checkpoints\checkpoint-25858\trainer_state.json"
OUTPUT_DOCX = r"D:\项目流程\BERT立场微调训练报告.docx"
CHART_DIR = r"D:\项目流程\charts"
os.makedirs(CHART_DIR, exist_ok=True)

# 中文字体
plt.rcParams['font.family'] = 'Microsoft YaHei'
plt.rcParams['axes.unicode_minus'] = False

# ===== 配色 =====
BLUE = '#4A90D9'
ORANGE = '#F5A623'
GREEN = '#7ED321'
RED = '#D0021B'
PURPLE = '#9013FE'
GRAY = '#9B9B9B'
COLORS_CLASS = [RED, GREEN, BLUE]  # tough, cooperate, neutral

# ===== 读取数据 =====
with open(TRAINER_STATE, 'r', encoding='utf-8') as f:
    state = json.load(f)

best_metric = state['best_metric']
best_step = state['best_global_step']
best_checkpoint = state['best_model_checkpoint']

# 提取 eval 记录
eval_records = []
for entry in state['log_history']:
    if 'eval_macro_f1' in entry:
        eval_records.append(entry)

epochs = [r['epoch'] for r in eval_records]
accs = [r['eval_accuracy'] for r in eval_records]
macro_f1s = [r['eval_macro_f1'] for r in eval_records]
eval_losses = [r['eval_loss'] for r in eval_records]

# 提取各类别 F1
class_f1 = {}
for cls_name in ['tough', 'cooperate', 'neutral']:
    key = f'eval_f1_{cls_name}'
    class_f1[cls_name] = [r[key] for r in eval_records]

# 提取 training loss
train_losses = []
train_steps = []
for entry in state['log_history']:
    if 'loss' in entry and 'eval_macro_f1' not in entry:
        train_steps.append(entry['step'])
        train_losses.append(entry['loss'])

# 计算训练时长
total_step_time = 0
# 从 log 估算

print(f"Epochs: {len(eval_records)}")
print(f"Best Macro F1: {best_metric:.4f} at step {best_step}")
print(f"Best checkpoint: {best_checkpoint}")

# =====================================================
# 图表 1：Macro F1 + Accuracy 趋势
# =====================================================
fig, ax1 = plt.subplots(figsize=(9, 5))

epoch_labels = [f'E{int(e)}' if e == int(e) else '' for e in epochs]
x = range(1, len(epochs) + 1)

line1, = ax1.plot(x, macro_f1s, 'o-', color=BLUE, linewidth=2.2, markersize=8, zorder=5, label='Macro F1')
ax1.set_ylabel('Macro F1', fontsize=13, color=BLUE)
ax1.set_ylim(0.5, 0.7)
ax1.tick_params(axis='y', labelcolor=BLUE)

# 标注最佳点
best_idx = macro_f1s.index(max(macro_f1s))
ax1.annotate(f'Best: {macro_f1s[best_idx]:.4f}',
             xy=(best_idx + 1, macro_f1s[best_idx]),
             xytext=(best_idx + 1 + 0.3, macro_f1s[best_idx] + 0.01),
             arrowprops=dict(arrowstyle='->', color=RED, lw=1.5),
             fontsize=11, color=RED, fontweight='bold')

ax2 = ax1.twinx()
line2, = ax2.plot(x, accs, 's--', color=ORANGE, linewidth=1.8, markersize=7, label='Accuracy')
ax2.set_ylabel('Accuracy', fontsize=13, color=ORANGE)
ax2.set_ylim(0.5, 0.7)
ax2.tick_params(axis='y', labelcolor=ORANGE)

ax1.set_xticks(x)
ax1.set_xticklabels(epoch_labels)
ax1.set_xlabel('Epoch', fontsize=13)

lines = [line1, line2]
labels = [l.get_label() for l in lines]
ax1.legend(lines, labels, loc='lower right', fontsize=11, framealpha=0.8)

plt.title('Macro F1 & Accuracy per Epoch', fontsize=15, fontweight='bold', pad=15)
ax1.grid(True, alpha=0.3, linestyle='--')
plt.tight_layout()
plt.savefig(os.path.join(CHART_DIR, '01_f1_accuracy.png'), dpi=200, bbox_inches='tight')
plt.close()
print("[图表] F1/Accuracy 趋势图已生成")

# =====================================================
# 图表 2：Training Loss + Eval Loss
# =====================================================
fig, (ax_top, ax_bot) = plt.subplots(2, 1, figsize=(10, 7), gridspec_kw={'height_ratios': [1, 1]})

# 上图：Training loss
ax_top.plot(train_steps, train_losses, color=BLUE, linewidth=0.6, alpha=0.7)
# 平滑曲线
if len(train_losses) > 20:
    window = max(len(train_losses) // 50, 5)
    smoothed = np.convolve(train_losses, np.ones(window)/window, mode='valid')
    smoothed_steps = train_steps[window//2:window//2+len(smoothed)]
    ax_top.plot(smoothed_steps, smoothed, color=RED, linewidth=2, label=f'Smoothed (window={window})')
ax_top.set_ylabel('Training Loss', fontsize=12)
ax_top.legend(fontsize=10)
ax_top.grid(True, alpha=0.3, linestyle='--')
ax_top.set_title('Training & Evaluation Loss', fontsize=15, fontweight='bold', pad=10)

# 下图：Eval loss
ax_bot.plot(x, eval_losses, 'D-', color=PURPLE, linewidth=2.2, markersize=8, zorder=5)
for i, (xi, yi) in enumerate(zip(x, eval_losses)):
    ax_bot.annotate(f'{yi:.3f}', (xi, yi), textcoords="offset points", xytext=(0, 12),
                    ha='center', fontsize=9, color=PURPLE)
ax_bot.set_xticks(x)
ax_bot.set_xticklabels(epoch_labels)
ax_bot.set_xlabel('Epoch', fontsize=12)
ax_bot.set_ylabel('Eval Loss', fontsize=12, color=PURPLE)
ax_bot.tick_params(axis='y', labelcolor=PURPLE)
ax_bot.grid(True, alpha=0.3, linestyle='--')

plt.tight_layout()
plt.savefig(os.path.join(CHART_DIR, '02_loss.png'), dpi=200, bbox_inches='tight')
plt.close()
print("[图表] Loss 曲线图已生成")

# =====================================================
# 图表 3：各类别 F1 柱状图（最后一轮）
# =====================================================
fig, ax = plt.subplots(figsize=(7, 5))
cls_names = ['Tough', 'Cooperate', 'Neutral']
last_f1 = [class_f1[k][-1] for k in ['tough', 'cooperate', 'neutral']]
best_f1_cls = [class_f1[k][best_idx] for k in ['tough', 'cooperate', 'neutral']]

x_pos = np.arange(len(cls_names))
width = 0.35
bars1 = ax.bar(x_pos - width/2, best_f1_cls, width, color=COLORS_CLASS, alpha=0.85, edgecolor='white', linewidth=0.8, label=f'Best (Epoch {best_idx+1})')
bars2 = ax.bar(x_pos + width/2, last_f1, width, color=COLORS_CLASS, alpha=0.35, edgecolor=GRAY, linewidth=0.8, linestyle='--', label=f'Last (Epoch {len(eval_records)})')

for bar in bars1:
    ax.text(bar.get_x() + bar.get_width()/2., bar.get_height() + 0.005, f'{bar.get_height():.3f}',
            ha='center', va='bottom', fontsize=11, fontweight='bold')
for bar in bars2:
    ax.text(bar.get_x() + bar.get_width()/2., bar.get_height() + 0.005, f'{bar.get_height():.3f}',
            ha='center', va='bottom', fontsize=10, color=GRAY)

ax.set_xticks(x_pos)
ax.set_xticklabels(cls_names, fontsize=13)
ax.set_ylabel('F1 Score', fontsize=13)
ax.set_ylim(0, 0.85)
ax.set_title('Per-Class F1 Score', fontsize=15, fontweight='bold')
ax.legend(fontsize=11)
ax.grid(True, alpha=0.2, axis='y', linestyle='--')
plt.tight_layout()
plt.savefig(os.path.join(CHART_DIR, '03_class_f1.png'), dpi=200, bbox_inches='tight')
plt.close()
print("[图表] 各类别 F1 柱状图已生成")

# =====================================================
# 图表 4：各类别 F1 趋势
# =====================================================
fig, ax = plt.subplots(figsize=(9, 5))
for i, cls in enumerate(['tough', 'cooperate', 'neutral']):
    ax.plot(x, class_f1[cls], 'o-', color=COLORS_CLASS[i], linewidth=2, markersize=7, label=cls.capitalize())

ax.set_xticks(x)
ax.set_xticklabels(epoch_labels)
ax.set_xlabel('Epoch', fontsize=13)
ax.set_ylabel('F1 Score', fontsize=13)
ax.set_title('Per-Class F1 Trend', fontsize=15, fontweight='bold')
ax.legend(fontsize=12, loc='lower right')
ax.grid(True, alpha=0.3, linestyle='--')
ax.set_ylim(0.3, 0.9)
plt.tight_layout()
plt.savefig(os.path.join(CHART_DIR, '04_class_trend.png'), dpi=200, bbox_inches='tight')
plt.close()
print("[图表] 各类别 F1 趋势图已生成")

# =====================================================
# 图表 5：学习率调度
# =====================================================
fig, ax = plt.subplots(figsize=(9, 3.5))
lr_records = [(e['step'], e['learning_rate']) for e in state['log_history'] if 'learning_rate' in e and 'eval_macro_f1' not in e]
lr_steps, lr_vals = zip(*lr_records)
ax.plot(lr_steps, lr_vals, color=BLUE, linewidth=1.5)
ax.set_xlabel('Step', fontsize=12)
ax.set_ylabel('Learning Rate', fontsize=12)
ax.set_title('Learning Rate Schedule (Linear Warmup + Linear Decay)', fontsize=13, fontweight='bold')
ax.grid(True, alpha=0.3, linestyle='--')
ax.ticklabel_format(style='scientific', axis='y', scilimits=(0, 0))
plt.tight_layout()
plt.savefig(os.path.join(CHART_DIR, '05_lr_schedule.png'), dpi=200, bbox_inches='tight')
plt.close()
print("[图表] 学习率调度图已生成")

# =====================================================
# 生成 Word 报告
# =====================================================
doc = Document()

# 页面设置
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# ---- 封面标题 ----
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('BERT 三分类立场微调\n训练报告')
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run(f'EU-China Corpus · Tough / Cooperate / Neutral')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

date_para = doc.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_para.add_run(f'生成时间：{datetime.now().strftime("%Y-%m-%d %H:%M")}')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

# ---- 目录函数 ----
def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    return h

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers), style='Light Grid Accent 1')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = str(h)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(10)
    # Data
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()
    return table

def add_image(path, width_inches=5.8):
    if os.path.exists(path):
        img_para = doc.add_paragraph()
        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_para.add_run().add_picture(path, width=Inches(width_inches))
        doc.add_paragraph()

# ============ 1. 实验概述 ============
add_heading('1. 实验概述', level=1)
doc.add_paragraph(
    '本实验旨在对 EU-China 政策语料进行三分类立场识别，使用 DistilBERT 预训练模型进行微调，'
    '将文本分为 "Tough（强硬）"、"Cooperate（合作）"、"Neutral（中立）" 三类。'
    '训练采用加权交叉熵损失以缓解类别不均衡问题，并通过早停策略（Early Stopping）'
    '防止过拟合。'
)

# ============ 2. 模型与超参数 ============
add_heading('2. 模型与超参数配置', level=1)

add_heading('2.1 模型架构', level=2)
add_table(
    ['参数', '值'],
    [
        ['基座模型', 'DistilBERT (distilbert-base-uncased)'],
        ['领域预训练', 'DAPT (Domain-Adaptive Pretraining)'],
        ['层数', '6 层 Transformer'],
        ['隐藏维度', '768'],
        ['注意力头数', '12'],
        ['总参数量', '66,955,011'],
        ['可训练参数', '66,955,011'],
        ['分类头', '3 标签 (Tough / Cooperate / Neutral)'],
    ]
)

add_heading('2.2 训练超参数', level=2)
add_table(
    ['参数', '值', '说明'],
    [
        ['最大序列长度', '128', '前128 token含核心立场信号'],
        ['学习率', '2e-5', 'AdamW 优化器'],
        ['Batch Size', '8', '受限于 MX450 2GB 显存'],
        ['梯度累积步数', '2', '等效 batch = 8×2 = 16'],
        ['最大训练轮次', '10', '—'],
        ['早停容忍', '2 轮', 'Macro F1 连续2轮不提升即停止'],
        ['最优指标', 'Macro F1', '用于早停和最优模型选择'],
        ['Weight Decay', '0.01', 'L2 正则化'],
        ['Warmup Ratio', '0.06', '线性预热'],
        ['学习率调度', 'Linear Decay', '线性衰减至 0'],
        ['最大梯度范数', '1.0', '梯度裁剪'],
        ['混合精度', 'FP16', 'NVIDIA MX450 支持'],
        ['随机种子', '42', '保证可复现'],
    ]
)

add_heading('2.3 损失函数', level=2)
doc.add_paragraph(
    '采用带类别权重的 CrossEntropyLoss（Balanced 策略），权重公式为：\n'
    '    weight[c] = n_total / (n_classes × n_class_samples)\n\n'
    '该策略使样本数较少的 "Tough" 和 "Neutral" 类别获得更高权重，'
    '抑制 "Cooperate" 类别样本过多导致的偏向。'
)

add_heading('2.4 数据集', level=2)
add_table(
    ['数据集', '样本数', '比例', '用途'],
    [
        ['训练集 (Train)', '~99,000', '~70%', '模型参数更新'],
        ['验证集 (Val)', '~19,000', '~14%', '早停判断 & 最优模型选择'],
        ['测试集 (Test)', '~23,000', '~16%', '最终评估（待进行）'],
    ]
)

# ============ 3. 训练结果 ============
add_heading('3. 训练结果', level=1)

add_heading('3.1 整体指标', level=2)
add_table(
    ['Epoch', 'Accuracy', 'Macro F1', 'Eval Loss', '备注'],
    [
        [f'{int(e)}', f'{accs[i]:.4f}', f'{macro_f1s[i]:.4f}', f'{eval_losses[i]:.4f}',
         '★ 最佳' if i == best_idx else ('早停参考' if i >= best_idx + 2 else '')]
        for i, e in enumerate(epochs)
    ]
)

add_heading('3.2 各类别 F1（最佳轮 Epoch {}）'.format(best_idx + 1), level=2)
add_table(
    ['类别', 'Precision', 'Recall', 'F1 Score', 'Support'],
    [
        ['Tough', f'{eval_records[best_idx].get("eval_precision_tough", "N/A")}',
         f'{eval_records[best_idx].get("eval_recall_tough", "N/A")}',
         f'{class_f1["tough"][best_idx]:.4f}', '—'],
        ['Cooperate', f'{eval_records[best_idx].get("eval_precision_cooperate", "N/A")}',
         f'{eval_records[best_idx].get("eval_recall_cooperate", "N/A")}',
         f'{class_f1["cooperate"][best_idx]:.4f}', '—'],
        ['Neutral', f'{eval_records[best_idx].get("eval_precision_neutral", "N/A")}',
         f'{eval_records[best_idx].get("eval_recall_neutral", "N/A")}',
         f'{class_f1["neutral"][best_idx]:.4f}', '—'],
    ]
)

add_heading('3.3 最优模型信息', level=2)
add_table(
    ['项目', '值'],
    [
        ['最优 Macro F1', f'{best_metric:.4f}'],
        ['最优 Global Step', str(best_step)],
        ['最优 Epoch', str(best_idx + 1)],
        ['模型路径', best_checkpoint],
        ['输出目录', r'D:\项目流程\bert_stance_model'],
    ]
)

# ============ 4. 图表 ============
add_heading('4. 训练曲线可视化', level=1)

add_heading('4.1 Macro F1 与 Accuracy 趋势', level=2)
doc.add_paragraph(
    '下图展示了各轮次的 Macro F1 和 Accuracy 变化。红色箭头标注了最佳 F1 所在轮次（Epoch {}）。'
    '可以看到 F1 在 Epoch 6 达到峰值后回落，说明此后模型开始过拟合。'.format(best_idx + 1)
)
add_image(os.path.join(CHART_DIR, '01_f1_accuracy.png'))

add_heading('4.2 损失曲线', level=2)
doc.add_paragraph(
    '上图：训练损失（Training Loss）持续下降，红线为平滑曲线。\n'
    '下图：验证损失（Eval Loss）从 Epoch 3 开始持续上升，表明模型在训练集上继续优化，'
    '但在验证集上的泛化性能已经下降——这是典型的过拟合信号。'
)
add_image(os.path.join(CHART_DIR, '02_loss.png'))

add_heading('4.3 各类别 F1 对比', level=2)
doc.add_paragraph(
    '对比最佳轮（Epoch {}）和最后一轮（Epoch {}）的各类别 F1 得分。'
    '"Cooperate" 类别表现最好，"Neutral" 类别因样本较多也较稳定，'
    '"Tough" 类别因样本少而 F1 相对偏低。'.format(best_idx + 1, len(eval_records))
)
add_image(os.path.join(CHART_DIR, '03_class_f1.png'))

add_heading('4.4 各类别 F1 趋势', level=2)
doc.add_paragraph('下图展示三类 F1 在各轮次中的变化趋势。')
add_image(os.path.join(CHART_DIR, '04_class_trend.png'))

add_heading('4.5 学习率调度', level=2)
doc.add_paragraph('采用 Linear Warmup + Linear Decay 策略，前 6% 步数线性预热后线性衰减至零。')
add_image(os.path.join(CHART_DIR, '05_lr_schedule.png'))

# ============ 5. 结论 ============
add_heading('5. 实验结论', level=1)

# 找出各类别 F1
cls_best = {k: class_f1[k][best_idx] for k in ['tough', 'cooperate', 'neutral']}
best_cls = max(cls_best, key=cls_best.get)
worst_cls = min(cls_best, key=cls_best.get)

doc.add_paragraph(
    f'1. 最优模型在 Epoch {best_idx + 1} 取得，Macro F1 = {best_metric:.4f}，Accuracy = {accs[best_idx]:.4f}。\n\n'
    f'2. 三类立场识别表现：{cls_best["tough"]:.3f} (Tough) / '
    f'{cls_best["cooperate"]:.3f} (Cooperate) / {cls_best["neutral"]:.3f} (Neutral)。\n'
    f'   "{best_cls.capitalize()}" 类别表现最优（F1={cls_best[best_cls]:.4f}），'
    f'"{worst_cls.capitalize()}" 类别相对较弱（F1={cls_best[worst_cls]:.4f}）。\n\n'
    f'3. 早停（patience=2）未触发，但 Epoch {best_idx + 1} 后 Eval Loss 持续上升、F1 不再提升，'
    f'确认过拟合已发生。建议以 Epoch {best_idx + 1} 模型为最终模型。\n\n'
    f'4. 后续改进方向：(a) 增大 Tough 类样本量或使用数据增强；'
    f'(b) 尝试更大的 BERT 模型（如 BERT-base）；'
    f'(c) 调整早停 patience 为 1 以减少无效训练。'
)

add_heading('5.1 训练环境', level=2)
add_table(
    ['项目', '值'],
    [
        ['GPU', 'NVIDIA GeForce MX450 (2GB)'],
        ['CPU', 'Intel Core i7-11370H @ 3.30GHz'],
        ['内存', '16 GB'],
        ['操作系统', 'Windows 10'],
        ['Python', '3.13'],
        ['深度学习框架', 'PyTorch + Hugging Face Transformers'],
        ['单轮耗时', '约 2-3 小时'],
        ['总训练轮数', str(len(eval_records))],
    ]
)

# 保存
doc.save(OUTPUT_DOCX)
print(f"\n{'='*60}")
print(f"  报告已生成：{OUTPUT_DOCX}")
print(f"{'='*60}")
